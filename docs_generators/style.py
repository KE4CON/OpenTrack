# -*- coding: utf-8 -*-
"""
Shared styling + building blocks for the APRS Command User Manual.

The manual's *source of truth* is this build pipeline (style.py + build.py + the
chapter modules). Running build.py regenerates docs/published/USER_MANUAL.docx from scratch.

IMPORTANT WORKFLOW NOTE
-----------------------
Regenerating overwrites the .docx. Screenshots are added by hand in Word only
AFTER every chapter's text is complete — until then, regenerating is safe.
Once screenshots are in, do NOT regenerate; make further changes in Word, and
ship future changes as dated amendment supplements (see the amendment model).
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT as WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette --------------------------------------------------------------
INK      = RGBColor(0x22, 0x2B, 0x38)
ACCENT   = RGBColor(0x1F, 0x3A, 0x5F)   # navy (primary)
ACCENT2  = RGBColor(0x2E, 0x6D, 0xA4)   # medium blue (section heads)
MUTED    = RGBColor(0x64, 0x74, 0x8B)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_HEX = "1F3A5F"
# navy + gold identity (matches the FieldCommand-style design language)
NAVY      = ACCENT
NAVY_HEX  = "1F3A5F"
NAVY_DEEP_HEX = "17314F"                 # slightly darker band
GOLD      = RGBColor(0xD9, 0xA5, 0x21)   # gold accent (marquee rules)
GOLD_HEX  = "D9A521"
PANEL_HEX = "EAF0F6"                      # pale banner fill (chapter title block)
COVER_MUTE = RGBColor(0xC7, 0xD6, 0xE8)   # muted light text on navy
COVER_FAINT = RGBColor(0x9F, 0xB3, 0xCC)

CALLOUTS = {
    "important": ("FFF4CC", "E0A800", RGBColor(0x8A, 0x6D, 0x00)),
    "note":      ("E7F0FA", "2E6DA4", ACCENT),
    "tip":       ("E8F5E9", "43A047", RGBColor(0x2E, 0x7D, 0x32)),
    "warning":   ("FDECEA", "D9534F", RGBColor(0xA3, 0x2A, 0x2A)),
}
SHOT_FILL, SHOT_BORDER = "F1F5F9", "94A3B8"

# ---- low-level xml helpers ------------------------------------------------
def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _borders(cell, sides):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in sides:
            s = sides[edge]
            el = OxmlElement('w:' + edge)
            el.set(qn('w:val'), s.get('val', 'single'))
            el.set(qn('w:sz'), str(s.get('sz', 6)))
            el.set(qn('w:space'), str(s.get('space', 0)))
            el.set(qn('w:color'), s.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def _margins(cell, top=110, start=170, bottom=110, end=170):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        el = OxmlElement('w:' + m); el.set(qn('w:w'), str(v)); el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def par_border(p, edge, sz=4, color='CBD5E1', val='single', space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    el = OxmlElement('w:' + edge)
    el.set(qn('w:val'), val); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), str(space)); el.set(qn('w:color'), color)
    pbdr.append(el)

def _field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); run._r.append(b)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; run._r.append(it)
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate'); run._r.append(sep)
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); run._r.append(e)
    return run

def _set_row_height(row, inches, rule='atLeast'):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), str(int(inches * 1440))); h.set(qn('w:hRule'), rule)
    trPr.append(h)

def _no_borders(cell):
    _borders(cell, {e: {'val': 'nil'} for e in ('top', 'left', 'bottom', 'right')})

# ---- run helpers ----------------------------------------------------------
def add_runs(p, parts):
    """parts: string, or list of str / (text, 'bi') tuples."""
    if not isinstance(parts, list):
        parts = [parts]
    for part in parts:
        if isinstance(part, tuple):
            t, st = part; r = p.add_run(t)
            if 'b' in st: r.bold = True
            if 'i' in st: r.italic = True
            if 'c' in st:  # code / monospace-ish
                r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        else:
            p.add_run(part)

# ---- document setup -------------------------------------------------------
def new_document(header_title="APRS Command — User Manual",
                 header_sub="Situational awareness for amateur radio  ·  Receive-first, transmit-safe",
                 footer_left="APRS Command  ·  Open-source (GPL v3)  ·  73 de KE4CON"):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
    sec.different_first_page_header_footer = True  # clean title page

    normal = doc.styles['Normal']
    normal.font.name = 'Segoe UI'; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.15; normal.paragraph_format.space_after = Pt(8)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI Semibold'; h1.font.size = Pt(19); h1.font.color.rgb = ACCENT; h1.font.bold = False
    h1.paragraph_format.space_before = Pt(18); h1.paragraph_format.space_after = Pt(6)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI Semibold'; h2.font.size = Pt(13.5); h2.font.color.rgb = ACCENT2; h2.font.bold = False
    h2.paragraph_format.space_before = Pt(14); h2.paragraph_format.space_after = Pt(4)

    # running header: centered title + subtitle, gold rule beneath (non-title pages)
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run(header_title)
    hr.font.name = 'Segoe UI Semibold'; hr.font.size = Pt(9.5); hr.font.color.rgb = NAVY
    hp.add_run().add_break()
    hs = hp.add_run(header_sub)
    hs.font.name = 'Segoe UI'; hs.font.size = Pt(8); hs.font.color.rgb = MUTED
    par_border(hp, 'bottom', sz=12, color=GOLD_HEX, space=6)

    # footer: navy rule, then attribution (left) + "Page X of Y" (right) via a 2-col table
    fp = sec.footer.paragraphs[0]; fp.paragraph_format.space_after = Pt(2)
    par_border(fp, 'bottom', sz=6, color=NAVY_HEX, space=2)
    ftbl = sec.footer.add_table(rows=1, cols=2, width=Inches(6.5))
    ftbl.autofit = False; ftbl.allow_autofit = False
    lc = ftbl.cell(0, 0); lc.width = Inches(4.5)
    rc = ftbl.cell(0, 1); rc.width = Inches(2.0)
    _no_borders(lc); _no_borders(rc)
    lp = lc.paragraphs[0]; lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run(footer_left)
    lr.font.name = 'Segoe UI'; lr.font.size = Pt(8); lr.font.color.rgb = MUTED
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rp.paragraph_format.space_after = Pt(0)
    r1 = rp.add_run("Page "); r1.font.name = 'Segoe UI'; r1.font.size = Pt(8); r1.font.color.rgb = MUTED
    pf = _field(rp, ' PAGE '); pf.font.size = Pt(8); pf.font.color.rgb = MUTED
    r2 = rp.add_run(" of "); r2.font.name = 'Segoe UI'; r2.font.size = Pt(8); r2.font.color.rgb = MUTED
    nf = _field(rp, ' NUMPAGES '); nf.font.size = Pt(8); nf.font.color.rgb = MUTED
    return doc

# ---- content blocks -------------------------------------------------------
def body(doc, parts):
    p = doc.add_paragraph(); add_runs(p, parts); return p

def h1(doc, text):
    return doc.add_heading(text, level=1)

def h2(doc, text):
    return doc.add_heading(text, level=2)

def step(doc, n, parts):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.38); pf.first_line_indent = Inches(-0.38); pf.space_after = Pt(5)
    pf.tab_stops.add_tab_stop(Inches(0.38), WD_TAB_ALIGNMENT.LEFT)
    r = p.add_run("%d.\t" % n); r.bold = True; r.font.color.rgb = ACCENT
    add_runs(p, parts)

def steps(doc, items):
    for i, it in enumerate(items, 1):
        step(doc, i, it)

def bullet(doc, parts):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.38); pf.first_line_indent = Inches(-0.20); pf.space_after = Pt(4)
    pf.tab_stops.add_tab_stop(Inches(0.38), WD_TAB_ALIGNMENT.LEFT)
    r = p.add_run(u"▪\t"); r.font.color.rgb = ACCENT2
    add_runs(p, parts)

def bullets(doc, items):
    for it in items:
        bullet(doc, it)

def _one_cell(doc, width=Inches(6.5)):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; tbl.allow_autofit = False
    cell = tbl.cell(0, 0); cell.width = width
    # Keep the whole box together — never let a screenshot, callout, or code block split across a
    # page break (image separated from its caption, or a half-box stranded at a page bottom). This
    # matters most once real screenshots replace the placeholders: the box moves to the next page as
    # a unit rather than straddling the break.
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))
    return tbl, cell

def callout(doc, kind, label, text):
    fill, border, labelcolor = CALLOUTS[kind]
    _, cell = _one_cell(doc)
    _shade(cell, fill)
    _borders(cell, {'left': {'sz': 30, 'color': border}, 'top': {'sz': 4, 'color': border},
                    'bottom': {'sz': 4, 'color': border}, 'right': {'sz': 4, 'color': border}})
    _margins(cell)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    rl = p.add_run(label + "   "); rl.bold = True; rl.font.color.rgb = labelcolor; rl.font.size = Pt(10.5)
    add_runs(p, text if isinstance(text, list) else [text])
    for r in p.runs[1:]:
        if r.font.color.rgb is None:
            r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def screenshot(doc, caption):
    _, cell = _one_cell(doc)
    _shade(cell, SHOT_FILL)
    dashed = {'val': 'dashed', 'sz': 8, 'color': SHOT_BORDER}
    _borders(cell, {'left': dashed, 'top': dashed, 'bottom': dashed, 'right': dashed})
    _margins(cell, top=220, bottom=220)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SCREENSHOT"); r.bold = True; r.font.color.rgb = MUTED; r.font.size = Pt(9)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(caption); r2.italic = True; r2.font.color.rgb = MUTED; r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def page_break(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

def code_block(doc, text):
    """A shaded monospace command block. `text` is one command (str) or several (list of str)."""
    _, cell = _one_cell(doc)
    _shade(cell, "F3F4F6")
    _borders(cell, {'left': {'sz': 18, 'color': 'B7C0CC'}, 'top': {'sz': 2, 'color': 'CBD5E1'},
                    'bottom': {'sz': 2, 'color': 'CBD5E1'}, 'right': {'sz': 2, 'color': 'CBD5E1'}})
    _margins(cell, top=80, bottom=80, start=150, end=150)
    lines = text if isinstance(text, list) else [text]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line); r.font.name = 'Consolas'; r.font.size = Pt(9.5); r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ---- cover -----------------------------------------------------------------
def cover(doc, kicker, big_title, subtitle, doc_kind, version, tagline, author, date_str):
    """Navy title panel with a gold marquee rule — the manual's front cover."""
    top = doc.add_paragraph(); top.paragraph_format.space_after = Pt(0)
    par_border(top, 'bottom', sz=26, color=GOLD_HEX, space=1)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; tbl.allow_autofit = False
    cell = tbl.cell(0, 0); cell.width = Inches(6.5)
    _shade(cell, NAVY_HEX); _no_borders(cell)
    _set_row_height(tbl.rows[0], 8.0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _margins(cell, top=260, bottom=260, start=320, end=320)

    def line(text, *, size, color, name='Segoe UI', bold=False, before=0, after=6, rule=False):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
        if rule:
            par_border(p, 'bottom', sz=12, color=GOLD_HEX, space=6)
        if text:
            r = p.add_run(text); r.font.name = name; r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold
        return p

    # first paragraph already exists in the cell; use it for the kicker
    p0 = cell.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER; p0.paragraph_format.space_after = Pt(2)
    kr = p0.add_run(kicker); kr.font.name = 'Segoe UI Semibold'; kr.font.size = Pt(11); kr.font.color.rgb = GOLD

    line(big_title, size=46, color=WHITE, name='Segoe UI Semibold', bold=True, before=46, after=4)
    line(subtitle + "   " + version, size=16, color=GOLD, name='Segoe UI Semibold', after=12)
    line(None, size=1, color=GOLD, rule=True, after=14)
    line(doc_kind, size=24, color=WHITE, name='Segoe UI Semibold', bold=True, after=6)
    line(tagline, size=10.5, color=COVER_MUTE, before=26, after=12)
    line(author, size=12, color=WHITE, name='Segoe UI Semibold', bold=True, before=0, after=2)
    line(date_str, size=10, color=COVER_FAINT, before=2, after=0)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# ---- chapter opener -------------------------------------------------------
def chapter_open(doc, number, title, subtitle, in_this_chapter=None):
    page_break(doc)
    # number block (navy square) + title banner (pale blue), gold rule under
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; tbl.allow_autofit = False
    numcell = tbl.cell(0, 0); numcell.width = Inches(0.95)
    titcell = tbl.cell(0, 1); titcell.width = Inches(5.55)
    _shade(numcell, NAVY_HEX); _shade(titcell, PANEL_HEX)
    _no_borders(numcell); _no_borders(titcell)
    numcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    titcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_row_height(tbl.rows[0], 0.72)
    _margins(numcell, top=40, bottom=40, start=40, end=40)
    _margins(titcell, top=60, bottom=60, start=220, end=140)

    np = numcell.paragraphs[0]; np.alignment = WD_ALIGN_PARAGRAPH.CENTER; np.paragraph_format.space_after = Pt(0)
    nr = np.add_run(str(number)); nr.font.name = 'Segoe UI Semibold'; nr.font.size = Pt(30); nr.font.color.rgb = WHITE; nr.bold = True

    t = titcell.paragraphs[0]; t.paragraph_format.space_after = Pt(0)
    tr = t.add_run(title); tr.font.name = 'Segoe UI Semibold'; tr.font.size = Pt(22); tr.font.color.rgb = NAVY; tr.bold = True
    # register as Heading 1 outline level so it shows in the TOC
    _set_outline_level(t, 0)
    _style_as_toc_entry(t, title, level=1)

    rule = doc.add_paragraph(); rule.paragraph_format.space_before = Pt(3); rule.paragraph_format.space_after = Pt(8)
    par_border(rule, 'bottom', sz=16, color=GOLD_HEX, space=6)

    if subtitle:
        s = doc.add_paragraph()
        sr = s.add_run(subtitle); sr.italic = True; sr.font.size = Pt(11.5); sr.font.color.rgb = MUTED
    if in_this_chapter:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        itc = doc.add_paragraph(); itc.paragraph_format.space_after = Pt(2)
        ir = itc.add_run("IN THIS CHAPTER"); ir.bold = True; ir.font.size = Pt(9)
        ir.font.color.rgb = MUTED; ir.font.name = 'Segoe UI Semibold'
        for item in in_this_chapter:
            b = doc.add_paragraph(); b.paragraph_format.left_indent = Inches(0.2); b.paragraph_format.space_after = Pt(1)
            rb = b.add_run(u"•  " + item); rb.font.size = Pt(10); rb.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def _set_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(level)); pPr.append(ol)

def _style_as_toc_entry(p, text, level=1):
    # Ensure the paragraph carries the built-in Heading style name so TOC \u picks it up,
    # while keeping our custom visual formatting via direct run formatting above.
    pPr = p._p.get_or_add_pPr()
    ps = pPr.find(qn('w:pStyle'))
    if ps is None:
        ps = OxmlElement('w:pStyle'); pPr.insert(0, ps)
    ps.set(qn('w:val'), 'Heading1')

def toc(doc):
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-2" \\h \\z \\u')

def section_title(doc, text):
    """A big front-matter title (Philosophy, How to Use, …) that also lists in the TOC at level 1."""
    h = doc.add_paragraph()
    r = h.add_run(text)
    r.font.name = 'Segoe UI Semibold'; r.font.size = Pt(24); r.font.color.rgb = ACCENT
    _set_outline_level(h, 0)
    _style_as_toc_entry(h, text, level=1)
    rule = doc.add_paragraph(); par_border(rule, 'bottom', sz=12, color=ACCENT_HEX, space=8)
    return h

def _table_light_borders(t, color="D7DEE7", sz=4):
    tblPr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def _keep_table_together(t):
    """Stop tables straddling a page break. Each row is marked no-split (never break a
    single row across pages); every row but the last is glued to the next row (keepNext)
    so a table that fits on one page moves to the next page as a whole unit rather than
    splitting. As a fallback for any table taller than a full page, the header row is set
    to repeat at the top of each page it continues onto."""
    rows = t.rows
    n = len(rows)
    for ri, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))
        if ri == 0 and trPr.find(qn('w:tblHeader')) is None:
            trPr.append(OxmlElement('w:tblHeader'))
        if ri < n - 1:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._p.get_or_add_pPr()
                    for tag in ('w:keepNext', 'w:keepLines'):
                        if pPr.find(qn(tag)) is None:
                            pPr.append(OxmlElement(tag))

def table(doc, headers, rows, widths):
    """Styled table: accent header row, zebra body rows, light borders.
    Each cell value is an add_runs-style string or list of runs."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.allow_autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = widths[i]
        _shade(c, ACCENT_HEX)
        _margins(c, top=70, bottom=70, start=120, end=120)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h); r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10); r.font.name = 'Segoe UI Semibold'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]; c.width = widths[i]
            if ri % 2 == 1: _shade(c, "F3F6FA")
            _margins(c, top=60, bottom=60, start=120, end=120)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            add_runs(p, val)
            for rr in p.runs:
                if rr.font.size is None: rr.font.size = Pt(9.5)
    _table_light_borders(t)
    _keep_table_together(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---- JSON chapter renderer ------------------------------------------------
# Agents produce chapters as JSON (see build.py). This renders that JSON deterministically,
# so a malformed chapter can only affect itself — it can never break the Python build.

_INLINE = re.compile(r'\*\*(.+?)\*\*|`([^`]+?)`|__(.+?)__')

def parse_inline(text):
    """Turn a marked-up string into add_runs parts: **bold**, `code`, __italic__."""
    text = str(text)
    parts, pos = [], 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            parts.append(text[pos:m.start()])
        if m.group(1) is not None:   parts.append((m.group(1), 'b'))
        elif m.group(2) is not None: parts.append((m.group(2), 'c'))
        else:                        parts.append((m.group(3), 'i'))
        pos = m.end()
    if pos < len(text):
        parts.append(text[pos:])
    return parts if parts else [text]

_CALLOUT_KINDS = {'note', 'tip', 'important', 'warning'}

def render_chapter(doc, ch, number):
    """Render one chapter dict (title, subtitle, in_this_chapter, blocks[]) into the document."""
    chapter_open(doc, number, ch.get('title', 'Untitled'), ch.get('subtitle', ''), ch.get('in_this_chapter'))
    for blk in ch.get('blocks', []):
        if not isinstance(blk, dict):
            continue
        if 'h1' in blk:
            h1(doc, str(blk['h1']))
        elif 'h2' in blk:
            h2(doc, str(blk['h2']))
        elif 'p' in blk:
            body(doc, parse_inline(blk['p']))
        elif 'steps' in blk:
            steps(doc, [parse_inline(s) for s in blk['steps']])
        elif 'bullets' in blk:
            bullets(doc, [parse_inline(s) for s in blk['bullets']])
        elif 'callout' in blk:
            c = blk['callout'] if isinstance(blk['callout'], dict) else {}
            kind = c.get('kind', 'note')
            if kind not in _CALLOUT_KINDS:
                kind = 'note'
            callout(doc, kind, str(c.get('label', kind.upper())), parse_inline(c.get('text', '')))
        elif 'screenshot' in blk:
            screenshot(doc, str(blk['screenshot']))
        elif 'code' in blk:
            code_block(doc, blk['code'])
        elif 'table' in blk:
            t = blk['table'] if isinstance(blk['table'], dict) else {}
            headers = [str(h) for h in t.get('headers', [])]
            rows = [[parse_inline(cell) for cell in row] for row in t.get('rows', [])]
            ncol = max(1, len(headers))
            widths = t.get('widths')
            if widths and len(widths) == ncol:
                widths = [Inches(float(w)) for w in widths]
            else:
                widths = [Inches(6.5 / ncol)] * ncol
            if headers and all(len(r) == ncol for r in rows):
                table(doc, headers, rows, widths)
