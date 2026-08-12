"""
OpenTrack Programming Guide — builder.

Markdown is the living source of truth; a styled .docx (navy+gold, matching the User Manual and
Installation Guide) is generated from the same chapter JSON. Chapters are validated JSON under
./chapters/*.json, each of the form {"order": N, "title", "subtitle", "in_this_chapter"[], "blocks"[]}
(or nested under a "chapter" key). Blocks reuse the shared renderer (h1/h2/p/steps/bullets/callout/
code/table), so this book looks and reads like the other OpenTrack documents.

Run:  python guide_build.py     -> writes ../guides/OpenTrack_Programming_Guide.docx
                                     and ./PROGRAMMING_GUIDE.md
Env:  GUIDE_OUT overrides the .docx path (used when Word has the file locked).
"""
import os
import sys
import glob
import json
import re
import datetime

HERE = os.path.dirname(__file__)
sys.path.insert(0, os.path.abspath(os.path.join(HERE, "..", "..", "docs_generators")))
import style as S  # noqa: E402
from docx.shared import Pt  # noqa: E402


def load_chapters():
    out = []
    for path in sorted(glob.glob(os.path.join(HERE, "chapters", "*.json"))):
        with open(path, encoding="utf-8") as f:
            raw = json.load(f)
        ch = raw.get("chapter", raw)
        order = int(raw.get("order", ch.get("order", 999)))
        out.append((order, ch))
    out.sort(key=lambda x: x[0])
    return out


# ---- Markdown emitter (source of truth) -----------------------------------
def _md_inline(text):
    # JSON marks italic as __x__; Markdown reads that as bold, so convert to *x*. ** and ` pass through.
    return re.sub(r"__(.+?)__", r"*\1*", str(text))


def _md_blocks(blocks):
    lines = []
    for blk in blocks:
        if not isinstance(blk, dict):
            continue
        if "h1" in blk:
            lines.append(f"\n## {blk['h1']}\n")
        elif "h2" in blk:
            lines.append(f"\n### {blk['h2']}\n")
        elif "p" in blk:
            lines.append(_md_inline(blk["p"]) + "\n")
        elif "steps" in blk:
            lines += [f"{i}. {_md_inline(s)}" for i, s in enumerate(blk["steps"], 1)]
            lines.append("")
        elif "bullets" in blk:
            lines += [f"- {_md_inline(b)}" for b in blk["bullets"]]
            lines.append("")
        elif "callout" in blk:
            c = blk["callout"] if isinstance(blk["callout"], dict) else {}
            label = c.get("label", c.get("kind", "NOTE").upper())
            lines.append(f"> **{label}** — {_md_inline(c.get('text', ''))}\n")
        elif "code" in blk:
            code = blk["code"]
            code = "\n".join(code) if isinstance(code, list) else str(code)
            lines.append("```csharp\n" + code + "\n```\n")
        elif "table" in blk:
            t = blk["table"] if isinstance(blk["table"], dict) else {}
            headers = [str(h) for h in t.get("headers", [])]
            rows = t.get("rows", [])
            if headers:
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for r in rows:
                    lines.append("| " + " | ".join(_md_inline(c) for c in r) + " |")
                lines.append("")
    return "\n".join(lines)


def to_markdown(chapters):
    parts = [
        "# OpenTrack Programming Guide\n",
        "*How and why the code works — a maintainer's field guide, in plain language.*\n",
        f"*Generated {datetime.date.today().strftime('%B %d, %Y')} · Markdown is the living source of truth.*\n",
        "\n---\n",
    ]
    for number, (_order, ch) in enumerate(chapters, 1):
        parts.append(f"\n# {number}. {ch.get('title', 'Untitled')}\n")
        if ch.get("subtitle"):
            parts.append(f"*{_md_inline(ch['subtitle'])}*\n")
        parts.append(_md_blocks(ch.get("blocks", [])))
    return "\n".join(parts)


# ---- styled .docx (generated artifact) ------------------------------------
def build_docx(chapters):
    doc = S.new_document(
        header_title="OpenTrack — Programming Guide",
        header_sub="How & why the code works  ·  For developers & maintainers",
        footer_left="OpenTrack Programming Guide  ·  Open-source (AGPL v3)  ·  KE4CON",
    )
    S.cover(
        doc,
        kicker="OPENTRACK  ·  FOR DEVELOPERS & MAINTAINERS",
        big_title="OpenTrack",
        subtitle="Programming Guide",
        doc_kind="HOW & WHY THE CODE WORKS",
        version="v1.0",
        tagline="Architecture  ·  Every Subsystem  ·  The Reasoning Behind It  ·  Built to Outlive Its Author",
        author="James Rospopo  ·  KE4CON",
        date_str=datetime.date.today().strftime("%B %d, %Y"),
    )
    S.section_title(doc, "About This Book")
    S.body(doc, "This book explains how OpenTrack is built and, more importantly, why it was built that "
                "way — in language a newcomer can follow, with enough depth for a maintainer to work "
                "confidently. Every explanation is drawn from the real source code. Read it beside the "
                "code, not instead of it.")
    S.page_break(doc)
    h = doc.add_paragraph()
    hr = h.add_run("Contents")
    hr.font.name = 'Segoe UI Semibold'
    hr.font.size = Pt(24)
    hr.font.color.rgb = S.ACCENT
    rule = doc.add_paragraph()
    S.par_border(rule, 'bottom', sz=12, color=S.GOLD_HEX, space=8)
    note = doc.add_paragraph()
    nr = note.add_run("If the list below is blank or out of date, click it and press F9 (or right-click "
                      "-> Update Field) to rebuild it.")
    nr.italic = True
    nr.font.size = Pt(9.5)
    nr.font.color.rgb = S.MUTED
    S.toc(doc)
    for number, (_order, ch) in enumerate(chapters, 1):
        S.render_chapter(doc, ch, number)
    out = os.environ.get("GUIDE_OUT") or os.path.join(
        HERE, "..", "guides", "OpenTrack_Programming_Guide.docx")
    out = os.path.abspath(out)
    doc.save(out)
    return out


def main():
    chapters = load_chapters()
    if not chapters:
        print("No chapters found in ./chapters/*.json — nothing to build yet.")
        return
    md = to_markdown(chapters)
    md_path = os.path.abspath(os.path.join(HERE, "PROGRAMMING_GUIDE.md"))
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)
    docx_path = build_docx(chapters)
    print(f"OK — {len(chapters)} chapter(s)")
    print(f"  Markdown: {md_path}")
    print(f"  Word:     {docx_path}")


if __name__ == "__main__":
    main()
